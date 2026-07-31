import sys
import tempfile
import unittest
import json
from pathlib import Path
from unittest import mock

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

import paper_format as pf
import paper_release_gate as release_gate


class PaperFormatTests(unittest.TestCase):
    def _front_matter(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")
        pf.abstract_title(doc)
        pf.body(doc, "摘要正文")
        pf.keywords(doc, "优化；预测")
        return doc

    def test_reference_template_styles_are_kept_but_sample_body_is_cleared(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            source = Document()
            source.add_paragraph("模板示例正文，不应进入论文")
            source.save(template)

            doc = pf.new_document(contest="cumcm", template_path=template)

        self.assertNotIn("模板示例正文", "\n".join(p.text for p in doc.paragraphs))

    def test_official_fixed_template_content_can_be_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "official.docx"
            source = Document()
            source.add_paragraph("官方固定摘要页")
            source.save(template)

            doc = pf.new_document(
                contest="cumcm",
                template_path=template,
                preserve_template_content=True,
            )

        self.assertIn("官方固定摘要页", "\n".join(p.text for p in doc.paragraphs))

    def test_cumcm_structure_validator_requires_abstract_and_keywords(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")

        errors = pf.validate_paper_structure(doc, contest="cumcm")

        self.assertTrue(any("摘要" in error for error in errors))
        self.assertTrue(any("关键词" in error for error in errors))

    def test_complete_cumcm_front_matter_passes(self):
        doc = self._front_matter()

        errors = pf.validate_paper_structure(doc, contest="cumcm", quality_checks=False)

        self.assertEqual(errors, [])

    def test_unconfigured_quality_thresholds_are_not_invented(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(doc, contest="cumcm")

        self.assertEqual(issues, [])

    def test_explicit_quality_targets_report_content_formula_figure_table_and_page_gaps(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            quality_targets={
                "min_content_units": 100,
                "min_equations": 1,
                "min_figures": 1,
                "min_tables": 1,
                "body_pages": 10,
                "abstract_fill_min": 0.5,
            },
        )

        for expected in ("字词单位", "公式", "图", "表", "正文渲染页数", "摘要页"):
            self.assertTrue(any(expected in issue for issue in issues), expected)

    def test_table_caption_must_be_referenced_in_body(self):
        doc = self._front_matter()
        pf.body(doc, "正文没有引用下面的表格。")
        pf.three_line_table(doc, [["变量", "值"], ["x", "1"]])
        doc.add_paragraph("表1 参数结果")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("表1" in issue and "正文" in issue for issue in issues))

    def test_reference_list_and_body_citations_are_bidirectionally_checked(self):
        doc = self._front_matter()
        pf.body(doc, "已有研究支持该方法[1]，但错误引用了[3]。")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] A. Author. A useful paper.")
        doc.add_paragraph("[2] B. Author. An uncited paper.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("[3]" in issue and "参考文献表" in issue for issue in issues))
        self.assertTrue(any("[2]" in issue and "未在正文引用" in issue for issue in issues))

    def test_compound_reference_citations_are_recognized(self):
        doc = self._front_matter()
        pf.body(doc, "相关方法见文献[1, 2]及文献[3-4]。")
        doc.add_paragraph("参考文献")
        for number in range(1, 5):
            doc.add_paragraph(f"[{number}] Reference {number}.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertFalse(any("未在正文引用" in issue for issue in issues))

    def test_total_pages_do_not_trigger_the_body_page_limit(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=31,
            require_rendered_pages=False,
        )

        self.assertFalse(any("超过" in issue and "30" in issue for issue in issues))

    def test_rendered_body_pages_trigger_only_explicit_official_maximum(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=45,
            rendered_body_pages=31,
            hard_constraints={"max_body_pages": 30},
        )

        self.assertTrue(any("正文官方上限" in issue and "30" in issue for issue in issues))

    def test_electronic_paper_limit_is_file_size_not_total_pages(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=80,
            rendered_body_pages=20,
            electronic_file_size_bytes=21 * 1024 * 1024,
            hard_constraints={"max_file_size_mb": 20},
        )

        self.assertTrue(any("文件" in issue and "20MB" in issue for issue in issues))

    def test_abstract_fill_uses_rendered_ratio_instead_of_fixed_word_count(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_body_pages=28,
            abstract_fill_ratio=0.60,
            quality_targets={"abstract_fill_min": 0.85},
        )

        self.assertTrue(any("填充率" in issue and "85%" in issue for issue in issues))

    def test_near_full_abstract_and_exact_target_body_meet_quality_targets(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=50,
            rendered_body_pages=30,
            abstract_fill_ratio=0.90,
            quality_targets={
                "body_pages": 30,
                "body_page_tolerance": 0,
                "abstract_fill_min": 0.85,
                "abstract_fill_max": 0.96,
            },
        )

        self.assertFalse(any("页质量目标" in issue or "填充率" in issue for issue in issues))

    def test_safe_save_rejects_skill_root(self):
        doc = self._front_matter()

        with self.assertRaisesRegex(ValueError, "PROJECT_ROOT"):
            pf.save_document(doc, pf.SKILL_ROOT, contest="cumcm")

    def test_final_filename_cannot_be_saved_before_release_gate(self):
        doc = self._front_matter()
        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaisesRegex(ValueError, "paper_release_gate"):
                pf.save_document(doc, tmp, filename="完整论文.docx", contest="cumcm")

    def _release_manifest(self):
        evidence = {
            "problem_scope": "本文研究测试系统的方案设计与结果验证问题。",
            "q1_method_result": "问题一采用约束优化模型并得到可行方案。",
            "validation": "滚动扰动检验表明核心结论在测试范围内稳定。",
        }
        return {
            "schema_version": "2.1",
            "competition": "测试竞赛",
            "edition": "2026",
            "official_rules": {
                "source": "https://example.test/rules",
                "verified_date": "2026-07-29",
                "submission_format": "docx",
                "hard_constraints": {},
            },
            "template": {
                "authority": "official", "path": "template.docx", "sha256": "abc"
            },
            "renderer": {"name": "Microsoft Word", "version": "2026"},
            "quality_target": {
                "blocking": True,
                "body_pages": 30, "body_page_tolerance": 0,
                "abstract_fill_min": 0.85, "abstract_fill_max": 0.96,
            },
            "abstract_audit": {
                "required_items": ["problem_scope", "q1_method_result", "validation"],
                "items": {
                    item: {
                        "requirement": f"核验摘要中的 {item} 内容",
                        "status": "pass",
                        "evidence": [evidence[item]],
                    }
                    for item in ("problem_scope", "q1_method_result", "validation")
                },
            },
        }

    @staticmethod
    def _release_abstract(manifest):
        return " ".join(
            snippet
            for item in manifest["abstract_audit"]["items"].values()
            for snippet in item["evidence"]
        )

    def test_release_gate_rejects_short_body(self):
        manifest = self._release_manifest()
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 18, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("正文实测 18 页" in error for error in errors))

    def test_release_gate_rejects_overfilled_abstract(self):
        manifest = self._release_manifest()
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.98},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("98.0%" in error for error in errors))

    def test_nonblocking_quality_target_is_reported_as_warning(self):
        manifest = self._release_manifest()
        manifest["quality_target"]["blocking"] = False
        warnings = []
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 18, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
            warnings,
        )
        self.assertEqual(errors, [])
        self.assertTrue(any("正文实测 18 页" in warning for warning in warnings))

    def test_custom_author_year_references_can_disable_numeric_bijection(self):
        doc = pf.new_document(contest="mcm-icm")
        doc.add_paragraph("Summary")
        doc.add_paragraph("Summary text")
        doc.add_paragraph("References")
        doc.add_paragraph("Author, A. (2024). Example study.")

        issues = pf.validate_paper_structure(
            doc,
            contest="mcm-icm",
            require_reference_section=True,
            check_numeric_reference_bijection=False,
        )

        self.assertEqual(issues, [])

    def test_structure_and_abstract_checks_include_layout_tables(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        table.cell(0, 0).text = "Summary"
        table.cell(1, 0).text = "Scope evidence in table."
        table.cell(2, 0).text = "Keywords: test"
        issues = pf.validate_paper_structure(
            doc,
            contest="custom",
            required_markers=["Summary", "Keywords: test"],
            check_figure_numbering=False,
            check_table_numbering=False,
            check_numeric_reference_bijection=False,
            excluded_table_objects=1,
        )
        abstract = release_gate._abstract_text(doc, {
            "abstract_heading_pattern": "^Summary$",
            "abstract_end_pattern": "^Keywords:",
        })
        self.assertEqual(issues, [])
        self.assertEqual(abstract, "Scope evidence in table.")

    def test_release_gate_rejects_unknown_and_inverted_constraints(self):
        manifest = self._release_manifest()
        manifest["official_rules"]["hard_constraints"] = {
            "max_page": 30,
            "min_body_pages": 20,
            "max_body_pages": 10,
        }
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("未识别字段: max_page" in error for error in errors))
        self.assertTrue(any("min_body_pages" in error and "不能大于" in error for error in errors))

    def test_transactional_overwrite_restores_both_outputs_on_commit_failure(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_docx = root / "source.docx"
            source_pdf = root / "source.pdf"
            output_docx = root / "final.docx"
            output_pdf = root / "final.pdf"
            source_docx.write_bytes(b"new-docx")
            source_pdf.write_bytes(b"new-pdf")
            output_docx.write_bytes(b"old-docx")
            output_pdf.write_bytes(b"old-pdf")
            original_replace = Path.replace
            failed = False

            def guarded_replace(path, target):
                nonlocal failed
                target = Path(target)
                if not failed and path.name.endswith(".tmp") and target == output_pdf:
                    failed = True
                    raise OSError("simulated second-output failure")
                return original_replace(path, target)

            with mock.patch.object(Path, "replace", guarded_replace):
                with self.assertRaisesRegex(OSError, "simulated"):
                    release_gate._publish_files(
                        [(source_docx, output_docx), (source_pdf, output_pdf)],
                        overwrite=True,
                    )

            self.assertEqual(output_docx.read_bytes(), b"old-docx")
            self.assertEqual(output_pdf.read_bytes(), b"old-pdf")
            self.assertFalse(list(root.glob(".*.tmp")))
            self.assertFalse(list(root.glob(".*.bak")))

    def test_release_gate_rejects_unapproved_fallback_template(self):
        manifest = self._release_manifest()
        manifest["template"]["authority"] = "fallback"
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("备用模板" in error for error in errors))

    def test_release_gate_requires_semantic_evidence(self):
        manifest = self._release_manifest()
        manifest["abstract_audit"]["items"]["validation"]["evidence"] = ["不存在的验证证据完整片段"]
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.90},
            "本文研究测试系统的方案设计与结果验证问题。 问题一采用约束优化模型并得到可行方案。",
        )
        self.assertTrue(any("validation" in error for error in errors))

    def test_release_gate_uses_manifest_problem_count_instead_of_fixed_three(self):
        for count in (2, 4):
            manifest = self._release_manifest()
            problem_items = [f"q{index}_method_result" for index in range(1, count + 1)]
            required = ["problem_scope", *problem_items, "validation"]
            manifest["abstract_audit"] = {
                "required_items": required,
                "items": {
                    item: {
                        "requirement": f"核验动态必答项 {item}",
                        "status": "pass",
                        "evidence": [f"这是 {item} 对应的独立摘要原文证据片段。"],
                    }
                    for item in required
                },
            }
            errors = release_gate.evaluate_release(
                manifest,
                {"body_pages": 30, "abstract_fill_ratio": 0.90},
                self._release_abstract(manifest),
            )
            self.assertEqual(errors, [], count)

    def test_release_gate_rejects_reused_or_placeholder_abstract_evidence(self):
        manifest = self._release_manifest()
        shared = manifest["abstract_audit"]["items"]["problem_scope"]["evidence"]
        manifest["abstract_audit"]["items"]["validation"]["evidence"] = shared
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("复用了完全相同" in error for error in errors))
        manifest = self._release_manifest()
        manifest["abstract_audit"]["items"]["validation"]["evidence"] = ["替换为摘要证据"]
        errors = release_gate.evaluate_release(
            manifest,
            {"body_pages": 30, "abstract_fill_ratio": 0.90},
            self._release_abstract(manifest),
        )
        self.assertTrue(any("占位文本" in error for error in errors))

    def test_result_registry_requires_generated_by_and_common_fields(self):
        registry = {
            "answer": {
                "value": 1.2,
                "unit": "1",
                "precision": 2,
                "source": "results/output.json",
                "generator": "python solve.py",
            }
        }
        errors = release_gate.validate_result_registry(registry)
        self.assertTrue(any("generated_by" in error for error in errors))

    def test_result_registry_rejects_code_as_value_source_and_nonfinite_value(self):
        registry = {
            "answer": {
                "value": float("nan"),
                "unit": "1",
                "precision": 2,
                "source": "solve.py",
                "generated_by": "python solve.py",
            }
        }
        errors = release_gate.validate_result_registry(registry)
        self.assertTrue(any("NaN" in error for error in errors))
        self.assertTrue(any("source" in error for error in errors))

    def test_release_run_locks_all_artifacts_and_supports_no_appendix_layout(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            results = root / "results"
            rendered = root / "rendered"
            results.mkdir()
            rendered.mkdir()

            template = root / "template.docx"
            Document().save(template)
            candidate = root / "candidate.docx"
            doc = Document()
            doc.add_paragraph("Test Paper")
            doc.add_paragraph("Summary")
            doc.add_paragraph(
                "This paper defines the tested system scope and decision objective. "
                "For problem one, constrained optimization produces a feasible plan. "
                "Independent perturbation checks support the reported conclusion."
            )
            doc.add_paragraph("Keywords: test")
            doc.add_paragraph("Body text")
            doc.save(candidate)

            pdf_path = rendered / "candidate.pdf"
            pdf = release_gate.fitz.open()
            page = pdf.new_page()
            page.insert_text((72, 72), "Summary", fontsize=14)
            page.insert_text(
                (72, 110),
                "This paper defines the tested system scope and decision objective. "
                "For problem one, constrained optimization produces a feasible plan. "
                "Independent perturbation checks support the reported conclusion.",
            )
            page.insert_text((72, 150), "Keywords: test")
            page = pdf.new_page()
            page.insert_text((72, 72), "Body text")
            pdf.save(pdf_path)
            pdf.close()

            registry_path = results / "result_registry.json"
            (results / "output.json").write_text('{"answer": 1.2}', encoding="utf-8")
            registry_path.write_text(json.dumps({
                "q1.answer": {
                    "value": 1.2,
                    "unit": "1",
                    "precision": 2,
                    "source": "results/output.json",
                    "generated_by": "python solve.py",
                }
            }), encoding="utf-8")
            (results / "constraints.json").write_text("{}", encoding="utf-8")
            (results / "data_audit.json").write_text("{}", encoding="utf-8")
            (results / "objective_map.json").write_text("{}", encoding="utf-8")
            (results / "oracle.json").write_text("{}", encoding="utf-8")
            checks = {
                name: {"status": "na", "evidence": [], "reason": "测试不适用"}
                for name in release_gate.red_team_gate.REQUIRED_CHECKS
            }
            checks["cross_output_consistency"] = {
                "status": "pass",
                "evidence": ["results/result_registry.json"],
                "reason": "",
            }
            checks["data_integrity_and_leakage"] = {
                "status": "pass",
                "evidence": ["results/data_audit.json"],
                "reason": "",
            }
            checks["objective_metric_alignment"] = {
                "status": "pass",
                "evidence": ["results/objective_map.json"],
                "reason": "",
            }
            checks["implementation_and_oracle"] = {
                "status": "pass",
                "evidence": ["results/oracle.json"],
                "reason": "",
                "oracle_audit": {
                    "coverage_statement": "覆盖发布测试的唯一核心例程",
                    "oracles": [{
                        "id": "O1",
                        "type": "hand_calculation",
                        "target": "发布测试约束",
                        "discriminating_case": "手算零余量边界",
                        "expected": "余量为零",
                        "observed": "余量为零",
                        "status": "pass",
                        "evidence": ["results/oracle.json"],
                    }],
                },
            }
            audit_path = results / "red_team_audit.json"
            audit_path.write_text(json.dumps({
                "schema_version": "1.3",
                "project": "release test",
                "status": "pass",
                "applicable_checks": checks,
                "parameter_audit": {"status": "na", "reason": "无参数"},
                "claims": [{
                    "id": "C1",
                    "text": "test",
                    "strength": "feasible",
                    "status": "pass",
                    "evidence": ["results/constraints.json"],
                    "search_scope": "test",
                    "falsification": "constraint fails",
                    "proof": "",
                }],
                "findings": [],
            }), encoding="utf-8")

            candidate_hash = release_gate._sha256(candidate)
            manifest = {
                "schema_version": "2.1",
                "competition": "Generic Modeling Contest",
                "edition": "2026",
                "contest_profile": "custom",
                "official_rules": {
                    "source": "https://example.test/rules",
                    "verified_date": "2026-07-29",
                    "submission_format": "both",
                    "hard_constraints": {"max_body_pages": 2},
                },
                "template": {
                    "authority": "official",
                    "path": "template.docx",
                    "sha256": release_gate._sha256(template),
                },
                "renderer": {"name": "PyMuPDF test renderer", "version": "1"},
                "layout": {
                    "abstract_page": 1,
                    "abstract_heading_pattern": "^Summary$",
                    "abstract_end_pattern": "^Keywords:",
                    "body_start_page": 2,
                    "body_end_page": 2,
                    "appendix_required": False,
                },
                "structure": {"required_markers": ["Summary"]},
                "quality_target": {},
                "abstract_audit": {
                    "required_items": ["scope", "q1", "validation"],
                    "items": {
                        "scope": {
                            "requirement": "State the problem scope and decision objective",
                            "status": "pass",
                            "evidence": ["This paper defines the tested system scope and decision objective."],
                        },
                        "q1": {
                            "requirement": "State the method and result for problem one",
                            "status": "pass",
                            "evidence": ["For problem one, constrained optimization produces a feasible plan."],
                        },
                        "validation": {
                            "requirement": "State evidence validating the conclusion",
                            "status": "pass",
                            "evidence": ["Independent perturbation checks support the reported conclusion."],
                        },
                    },
                },
                "artifacts": {
                    "candidate": {
                        "path": "candidate.docx",
                        "sha256": candidate_hash,
                    },
                    "rendered_pdf": {
                        "path": "rendered/candidate.pdf",
                        "sha256": release_gate._sha256(pdf_path),
                        "source_candidate_sha256": candidate_hash,
                    },
                    "result_registry": {
                        "path": "results/result_registry.json",
                        "sha256": release_gate._sha256(registry_path),
                    },
                    "red_team_audit": {
                        "path": "results/red_team_audit.json",
                        "sha256": release_gate._sha256(audit_path),
                    },
                },
            }
            manifest_path = results / "paper_release_manifest.json"
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            output = root / "complete.docx"
            submission_pdf = root / "complete.pdf"

            result = release_gate.run(
                manifest_path,
                candidate,
                pdf_path,
                output,
                submission_pdf_output=submission_pdf,
            )

            self.assertEqual(result["status"], "pass", result)
            self.assertTrue(output.is_file())
            self.assertTrue(submission_pdf.is_file())
            self.assertEqual(result["submission_files"], [str(output), str(submission_pdf)])

            manifest["official_rules"]["submission_format"] = "pdf"
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            internal_docx = root / "pdf-only-internal.docx"
            pdf_only = root / "pdf-only-submit.pdf"
            pdf_result = release_gate.run(
                manifest_path,
                candidate,
                pdf_path,
                internal_docx,
                submission_pdf_output=pdf_only,
            )
            self.assertEqual(pdf_result["status"], "pass", pdf_result)
            self.assertEqual(pdf_result["submission_files"], [str(pdf_only)])
            self.assertTrue(internal_docx.is_file())
            self.assertTrue(pdf_only.is_file())

    def test_sha256_comparison_is_case_insensitive(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.bin"
            path.write_bytes(b"paper-template")
            self.assertTrue(release_gate._hash_matches(path, release_gate._sha256(path).upper()))

    def test_release_evidence_cannot_point_back_into_skill_root(self):
        with self.assertRaisesRegex(ValueError, "SKILL_ROOT"):
            release_gate._resolve_path(
                release_gate.SKILL_ROOT / "SKILL.md",
                release_gate.SKILL_ROOT.parent,
            )


if __name__ == "__main__":
    unittest.main()
